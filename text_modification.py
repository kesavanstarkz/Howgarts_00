from text_extraction import extract_entities
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------- FASTAPI IMPORTS ----------------
from fastapi import FastAPI
from fastapi.responses import FileResponse
from pydantic import BaseModel

# ---------------- FASTAPI APP ----------------
app = FastAPI(
    title="SOW Generator API",
    description="Generate Statement of Work (DOCX) from raw input text",
    version="1.0.0"
)

# ---------------- REQUEST MODEL ----------------
class SOWRequest(BaseModel):
    input_text: str


@app.get("/")
def read_root():
    return {"status": "online", "message": "SOW Generator API is running"}


# ---------------- DOCX GENERATOR ----------------
def create_docx_sow(data, output_path):
    """Creates a professionally formatted DOCX Statement of Work."""
    doc = Document()

    # --- Header Section ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("STATEMENT OF WORK")
    run.bold = True
    run.underline = True
    run.font.size = Pt(16)

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = confidential.add_run("CONFIDENTIALITY: HIGH")
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)

    doc.add_paragraph(
        "\nThe information contained in this document shall be deemed Confidential Information "
        "to both parties. It shall not be disclosed, duplicated, or used for any purpose other "
        "than that stated herein, in whole or in part, without prior written consent of the other party.\n"
    )

    # --- Summary Table ---
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    summary_data = [
        ("Project Name", data.get("project_name", "")),
        ("Name of Client", data.get("client_name", "")),
        ("Effective Date", data.get("effective_date", "")),
        ("Type", data.get("type", "")),
    ]

    for label, value in summary_data:
        row_cells = table.add_row().cells
        p0 = row_cells[0].paragraphs[0]
        run0 = p0.add_run(label)
        run0.bold = True
        
        row_cells[1].text = str(value) if value else "N/A"

    # Scope row
    scope_cells = table.add_row().cells
    p_scope = scope_cells[0].paragraphs[0]
    run_scope = p_scope.add_run("Scope & Timeline")
    run_scope.bold = True
    
    timeline = data.get("scope_timeline", [])
    if isinstance(timeline, list):
        scope_cells[1].text = "\n".join(timeline) if timeline else "TBD"
    else:
        scope_cells[1].text = str(timeline)

    doc.add_page_break()

    # --- Detailed Content ---
    doc.add_heading("STATEMENT OF WORK", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"This Statement of Work (“SOW”) is entered into as of {data.get('effective_date', 'N/A')} "
        f"(Effective Date) by and between {data.get('client_name', 'N/A')}, and AGILISIUM CONSULTING LLC."
    )

    doc.add_heading("1. BACKGROUND:", level=2)
    doc.add_paragraph(data.get("background", "Agilisium and the Client have entered into an agreement for services."))
    
    if data.get("scope_description") and data.get("scope_description") != "N/A":
        doc.add_paragraph(data.get("scope_description"))

    doc.add_heading("2. TERM OF ENGAGEMENT:", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Effective Date: {data.get('effective_date', 'N/A')}\n").bold = True
    p.add_run(f"Termination Date: {data.get('termination_date', '31st Dec’24')}").bold = True

    doc.add_heading("3. PROJECT MANAGERS / CONTACT:", level=2)
    doc.add_paragraph(f"Agilisium Project Contact: Manager Name {{name@agilisium.com}}")
    doc.add_paragraph(f"Client Project Contact: {data.get('client_contact', 'TBD')}")

    doc.add_heading("4. SCOPE OF SERVICE:", level=2)
    doc.add_paragraph("The project will follow the phases outlined below:")
    
    phases = data.get("phases", {})
    if isinstance(phases, dict):
        for phase_name, details in phases.items():
            doc.add_heading(phase_name.replace("_", " ").title(), level=3)
            doc.add_paragraph(str(details))
    else:
        doc.add_paragraph(str(phases))

    doc.add_heading("5. ARCHITECTURE:", level=2)
    doc.add_paragraph("[Architecture Diagram Placeholder]")
    
    tech_stack = data.get("tech_stack", [])
    if tech_stack:
        p_tech = doc.add_paragraph()
        run_tech = p_tech.add_run("Tech Stack: ")
        run_tech.bold = True
        p_tech.add_run(", ".join(tech_stack) if isinstance(tech_stack, list) else str(tech_stack))

    doc.save(output_path)
    return output_path


# ---------------- API ENDPOINT ----------------
@app.post("/generate-sow")
async def generate_sow_endpoint(request: SOWRequest):
    """
    Generate SOW DOCX from raw text
    """
    print(f"Received request for SOW generation...")
    extracted_data = extract_entities(request.input_text)
    print(f"Extracted data: {extracted_data}")

    output_file = "Statement_Of_Work_Generated.docx"
    create_docx_sow(extracted_data, output_file)
    
    print(f"SOW generated at {output_file}")

    return FileResponse(
        path=output_file,
        filename="Statement_Of_Work.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ---------------- RUNNER ----------------
if __name__ == "__main__":
    import uvicorn
    import sys

    # If arguments are passed, it might be a test run or a server start
    if len(sys.argv) > 1 and sys.argv[1] == "test":
        input_text = """
        Project Name: BOCE Advanced Information Portal
        Client: Alexion Pharmaceuticals, Inc.
        Date: 25th May 2024
        Type: Fixed Bid
        """
        extracted_data = extract_entities(input_text)
        create_docx_sow(extracted_data, "Test_SOW.docx")
        print("Test SOW generated.")
    else:
        print("Starting FastAPI server...")
        uvicorn.run(app, host="0.0.0.0", port=8000)
